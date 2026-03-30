"""
============================================================================
AI Procurement Readiness Tool - Word Report Generator
============================================================================
Version: 1.0
Date: 2026-03-30
Purpose: Generate professional Word documents from assessment results

LEARNING NOTE:
python-docx allows us to create .docx files programmatically.
We build documents by adding paragraphs, tables, and styles.

GOVERNANCE SIGNIFICANCE:
Reports provide human-readable summaries for procurement decisions.
Professional formatting ensures credibility with stakeholders.
============================================================================
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, List
import json


class AssessmentReportGenerator:
    """
    Generates professional Word documents from assessment data.
    
    WHAT THIS DOES:
    Takes assessment results from BigQuery and creates a formatted
    Word document that procurement officers can use for decision-making.
    
    STRUCTURE:
    1. Cover Page
    2. Executive Summary
    3. Readiness Scorecard
    4. Principle-by-Principle Analysis
    5. Recommendations
    6. Appendix
    """
    
    def __init__(self):
        """Initialize the document with default styles."""
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """
        Set up document-wide styles and formatting.
        
        LEARNING NOTE:
        python-docx uses "styles" to maintain consistent formatting.
        We define these once and reuse throughout the document.
        """
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
    
    def generate_report(
        self,
        assessment_data: Dict,
        output_filename: str = None
    ) -> str:
        """
        Generate complete assessment report.
        
        Args:
            assessment_data: Assessment record from BigQuery
            output_filename: Optional custom filename
            
        Returns:
            Path to generated Word document
            
        Example:
            >>> generator = AssessmentReportGenerator()
            >>> generator.generate_report(assessment_data)
            'assessment_ASM-2026-001.docx'
        """
        
        # Generate filename if not provided
        if output_filename is None:
            assessment_id = assessment_data['assessment_id']
            output_filename = f"assessment_{assessment_id}.docx"
        
        # Build report sections
        self._add_cover_page(assessment_data)
        self._add_page_break()
        
        self._add_executive_summary(assessment_data)
        self._add_page_break()
        
        self._add_readiness_scorecard(assessment_data)
        self._add_page_break()
        
        self._add_principle_analysis(assessment_data)
        self._add_page_break()
        
        self._add_recommendations(assessment_data)
        self._add_page_break()
        
        self._add_appendix(assessment_data)
        
        # Save document
        self.doc.save(output_filename)
        print(f"✓ Report generated: {output_filename}")
        
        return output_filename
    
    def _add_cover_page(self, data: Dict):
        """
        Create professional cover page.
        
        WHAT IT INCLUDES:
        - Title
        - Assessment ID
        - Vendor name
        - Date
        - Overall readiness score
        """
        # Title
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("AI PROCUREMENT READINESS\nASSESSMENT REPORT")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Assessment details
        details = [
            f"Assessment ID: {data['assessment_id']}",
            f"Vendor: {data['vendor_name']}",
            f"AI System: {data.get('system_name', 'N/A')}",
            f"Assessment Date: {self._format_date(data['assessment_date'])}",
            f"Assessor: {data['assessor_id']}",
            f"Agency: {data.get('assessor_agency', 'N/A')}"
        ]
        
        for detail in details:
            p = self.doc.add_paragraph(detail)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(14)
        
        # Spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Overall readiness score (big and bold)
        score_para = self.doc.add_paragraph()
        score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        score_text = f"Overall Readiness: {data['total_readiness_score']:.0%}"
        run = score_para.add_run(score_text)
        run.font.size = Pt(28)
        run.font.bold = True
        
        # Color code the score
        score = data['total_readiness_score']
        if score >= 0.8:
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif score >= 0.6:
            run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        else:
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
    
    def _add_executive_summary(self, data: Dict):
        """
        Add executive summary section.
        
        WHAT IT INCLUDES:
        - Overall assessment
        - Key strengths
        - Critical gaps
        - Recommendation
        """
        # Section heading
        self._add_heading("Executive Summary", level=1)
        
        # Overall assessment
        score = data['total_readiness_score']
        
        if score >= 0.8:
            status = "READY FOR PROCUREMENT"
            assessment = f"The vendor's AI system demonstrates strong readiness across most AI Verify principles with an overall score of {score:.0%}. The system meets or exceeds requirements in the majority of assessed areas."
        elif score >= 0.6:
            status = "CONDITIONALLY READY"
            assessment = f"The vendor's AI system shows moderate readiness with an overall score of {score:.0%}. While many principles are adequately addressed, several areas require improvement before full deployment."
        else:
            status = "NOT READY"
            assessment = f"The vendor's AI system demonstrates insufficient readiness with an overall score of {score:.0%}. Significant gaps exist across multiple principles that must be addressed before procurement can proceed."
        
        # Status
        status_para = self.doc.add_paragraph()
        run = status_para.add_run(f"Status: {status}")
        run.font.bold = True
        run.font.size = Pt(14)
        
        # Assessment text
        self.doc.add_paragraph(assessment)
        
        # Key strengths
        self._add_heading("Key Strengths", level=2)
        principle_scores = data['principle_scores']
        
        # Parse JSON if needed
        if isinstance(principle_scores, str):
            principle_scores = json.loads(principle_scores)
        
        strong_principles = [
            (name, score) for name, score in principle_scores.items()
            if score >= 0.8
        ]
        
        if strong_principles:
            for name, score in sorted(strong_principles, key=lambda x: x[1], reverse=True)[:3]:
                self.doc.add_paragraph(
                    f"• {name.replace('_', ' ')}: {score:.0%}",
                    style='List Bullet'
                )
        else:
            self.doc.add_paragraph("No principles scored above 80%.")
        
        # Critical gaps
        self._add_heading("Critical Gaps", level=2)
        weak_principles = [
            (name, score) for name, score in principle_scores.items()
            if score < 0.7
        ]
        
        if weak_principles:
            for name, score in sorted(weak_principles, key=lambda x: x[1])[:3]:
                self.doc.add_paragraph(
                    f"• {name.replace('_', ' ')}: {score:.0%} - Requires immediate attention",
                    style='List Bullet'
                )
        else:
            self.doc.add_paragraph("No critical gaps identified. All principles score above 70%.")
    
    def _add_readiness_scorecard(self, data: Dict):
        """
        Add visual scorecard table.
        
        WHAT IT INCLUDES:
        Table showing all 11 principles with scores and status indicators.
        """
        self._add_heading("Readiness Scorecard", level=1)
        
        # Parse principle scores
        principle_scores = data['principle_scores']
        if isinstance(principle_scores, str):
            principle_scores = json.loads(principle_scores)
        
        # Create table
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Principle'
        header_cells[1].text = 'Score'
        header_cells[2].text = 'Status'
        
        # Make header bold
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add data rows
        for principle_name, score in sorted(principle_scores.items()):
            row_cells = table.add_row().cells
            
            # Principle name (clean up formatting)
            clean_name = principle_name.replace('_', ' ').replace('P', 'P')
            row_cells[0].text = clean_name
            
            # Score
            row_cells[1].text = f"{score:.0%}"
            
            # Status with color coding
            if score >= 0.8:
                status = "✓ Ready"
                color = RGBColor(0, 128, 0)
            elif score >= 0.6:
                status = "⚠ Needs Work"
                color = RGBColor(255, 165, 0)
            else:
                status = "✗ Critical Gap"
                color = RGBColor(255, 0, 0)
            
            row_cells[2].text = status
            
            # Color the status
            for paragraph in row_cells[2].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = color
                    run.font.bold = True
        
        # Add overall score row
        row_cells = table.add_row().cells
        row_cells[0].text = "OVERALL READINESS"
        row_cells[1].text = f"{data['total_readiness_score']:.0%}"
        row_cells[2].text = "Average"
        
        # Make overall row bold
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
    
    def _add_principle_analysis(self, data: Dict):
        """
        Add detailed principle-by-principle analysis.
        
        WHAT IT INCLUDES:
        For each principle: score, questions answered, gaps identified.
        """
        self._add_heading("Principle-by-Principle Analysis", level=1)
        
        # Parse responses
        responses = data['responses']
        if isinstance(responses, str):
            responses = json.loads(responses)
        
        # Parse principle scores
        principle_scores = data['principle_scores']
        if isinstance(principle_scores, str):
            principle_scores = json.loads(principle_scores)
        
        # Group responses by principle
        principle_data = {}
        for response in responses:
            p_num = response['principle_number']
            p_name = response['principle_name']
            key = f"P{p_num}_{p_name.replace(' ', '')}"
            
            if key not in principle_data:
                principle_data[key] = []
            principle_data[key].append(response)
        
        # Analyze each principle
        for principle_key in sorted(principle_data.keys()):
            responses_list = principle_data[principle_key]
            score = principle_scores.get(principle_key, 0.0)
            
            # Principle heading
            clean_name = principle_key.replace('_', ' ')
            self._add_heading(f"{clean_name} - {score:.0%}", level=2)
            
            # Count response types
            have_count = sum(1 for r in responses_list if r['answer'] == 'Have')
            partial_count = sum(1 for r in responses_list if r['answer'] == 'Partial')
            gap_count = sum(1 for r in responses_list if r['answer'] == 'Gap')
            
            # Summary
            self.doc.add_paragraph(
                f"Total Checks: {len(responses_list)} | "
                f"Complete: {have_count} | "
                f"Partial: {partial_count} | "
                f"Gaps: {gap_count}"
            )
            
            # Gaps detail (if any)
            gaps = [r for r in responses_list if r['answer'] in ['Gap', 'Partial']]
            if gaps:
                self._add_heading("Areas for Improvement:", level=3)
                for gap in gaps[:5]:  # Show top 5 gaps
                    self.doc.add_paragraph(
                        f"• {gap['question_id']}: {gap['answer']} - {gap['question_text'][:100]}...",
                        style='List Bullet'
                    )
    
    def _add_recommendations(self, data: Dict):
        """
        Add procurement recommendations.
        
        WHAT IT INCLUDES:
        - Overall recommendation
        - Priority actions
        - Timeline suggestions
        """
        self._add_heading("Recommendations", level=1)
        
        score = data['total_readiness_score']
        
        # Overall recommendation
        if score >= 0.8:
            rec = "PROCEED WITH PROCUREMENT. The vendor demonstrates strong readiness. Monitor identified gaps during implementation."
        elif score >= 0.6:
            rec = "CONDITIONAL APPROVAL. Require vendor to address identified gaps before contract signing. Re-assess after improvements."
        else:
            rec = "DO NOT PROCEED. Significant gaps exist. Vendor must substantially improve AI governance before procurement can continue."
        
        self.doc.add_paragraph(rec).runs[0].font.bold = True
        
        # Priority actions
        self._add_heading("Priority Actions", level=2)
        
        principle_scores = data['principle_scores']
        if isinstance(principle_scores, str):
            principle_scores = json.loads(principle_scores)
        
        weak_areas = sorted(
            [(name, score) for name, score in principle_scores.items() if score < 0.7],
            key=lambda x: x[1]
        )
        
        if weak_areas:
            self.doc.add_paragraph("The vendor should prioritize improvements in:")
            for i, (name, score) in enumerate(weak_areas[:3], 1):
                self.doc.add_paragraph(
                    f"{i}. {name.replace('_', ' ')}: Currently {score:.0%}",
                    style='List Number'
                )
        else:
            self.doc.add_paragraph("Continue monitoring all principles during implementation.")
    
    def _add_appendix(self, data: Dict):
        """
        Add appendix with technical details.
        
        WHAT IT INCLUDES:
        - Assessment metadata
        - Scoring methodology
        - Evidence file list
        """
        self._add_heading("Appendix", level=1)
        
        # Assessment metadata
        self._add_heading("Assessment Metadata", level=2)
        
        metadata = [
            f"Assessment ID: {data['assessment_id']}",
            f"Vendor: {data['vendor_name']}",
            f"System: {data.get('system_name', 'N/A')}",
            f"Assessor: {data['assessor_id']}",
            f"Agency: {data.get('assessor_agency', 'N/A')}",
            f"Date: {self._format_date(data['assessment_date'])}",
            f"Status: {data['assessment_status']}",
            f"Contract Value: SGD {data.get('estimated_contract_value', 0):,.2f}"
        ]
        
        for item in metadata:
            self.doc.add_paragraph(item, style='List Bullet')
        
        # Scoring methodology
        self._add_heading("Scoring Methodology", level=2)
        self.doc.add_paragraph(
            "Readiness scores are calculated using the AI Verify Testing Framework formula:"
        )
        self.doc.add_paragraph(
            "Readiness = (Have + 0.5 × Partial) / Total Checks"
        )
        self.doc.add_paragraph(
            "• Have (1.0): Complete implementation with evidence\n"
            "• Partial (0.5): Incomplete implementation or documentation\n"
            "• Gap (0.0): Not implemented or no evidence"
        )
    
    def _add_heading(self, text: str, level: int = 1):
        """Add formatted heading."""
        heading = self.doc.add_heading(text, level=level)
        return heading
    
    def _add_page_break(self):
        """Add page break."""
        self.doc.add_page_break()
    
    def _format_date(self, date_str: str) -> str:
        """Format ISO date string to readable format."""
        try:
            dt = datetime.fromisoformat(date_str.replace('Z', '+00:00'))
            return dt.strftime('%d %B %Y')
        except:
            return date_str


# ============================================================================
# EXAMPLE USAGE
# ============================================================================

if __name__ == "__main__":
    print("=" * 60)
    print("WORD REPORT GENERATOR - TEST")
    print("=" * 60)
    
    # Sample assessment data (would normally come from BigQuery)
    sample_data = {
        "assessment_id": "ASM-2026-001",
        "vendor_name": "Example AI Corp",
        "system_name": "SmartDoc Analyzer",
        "assessment_date": "2026-03-30T10:00:00Z",
        "assessor_id": "officer_chan_123",
        "assessor_agency": "GovTech Singapore",
        "assessment_status": "approved",
        "total_readiness_score": 0.79,
        "principle_scores": {
            "P1_Transparency": 0.80,
            "P2_Explainability": 0.50,
            "P3_Reproducibility": 0.77,
            "P4_Safety": 0.83,
            "P5_Security": 0.82,
            "P6_Robustness": 0.71,
            "P7_Fairness": 0.70,
            "P8_DataGovernance": 1.00,
            "P9_Accountability": 0.82,
            "P10_HumanAgency": 0.75,
            "P11_InclusiveGrowth": 1.00
        },
        "responses": [
            {
                "question_id": "Q001",
                "principle_number": 1,
                "principle_name": "Transparency",
                "question_text": "Does the vendor align with PDPC guidelines?",
                "answer": "Have"
            },
            {
                "question_id": "Q002",
                "principle_number": 1,
                "principle_name": "Transparency",
                "question_text": "Does the vendor maintain transparency measures?",
                "answer": "Partial"
            }
        ],
        "estimated_contract_value": 500000.00
    }
    
    # Generate report
    generator = AssessmentReportGenerator()
    filename = generator.generate_report(sample_data, "test_report.docx")
    
    print(f"\n✓ Test report generated: {filename}")
    print("=" * 60)
